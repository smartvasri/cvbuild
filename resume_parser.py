"""
resume_parser.py
Extracts plain text from uploaded PDF or DOCX files.
"""

import io
from pdfminer.high_level import extract_text as _pdf_extract
from docx import Document


def extract_text_from_pdf(content: bytes) -> str:
    """Extract all text from a PDF file given its raw bytes."""
    pdf_stream = io.BytesIO(content)
    text = _pdf_extract(pdf_stream)
    return text.strip()


def extract_text_from_docx(content: bytes) -> str:
    """Extract all text from a DOCX file given its raw bytes."""
    docx_stream = io.BytesIO(content)
    doc = Document(docx_stream)

    lines = []

    # Extract paragraph text
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            lines.append(stripped)

    # Extract text from tables (skills tables, experience grids, etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped and stripped not in lines:
                    lines.append(stripped)

    return "\n".join(lines)


def extract_text(content: bytes, filename: str) -> str:
    """
    Dispatcher: choose the right extractor based on file extension.
    filename should be the original uploaded filename (lowercased).
    """
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are allowed.")
